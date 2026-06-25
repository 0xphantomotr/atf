import json
import re
import shlex
import subprocess
import uuid
from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory
from typing import Any

from pypdf import PdfReader
from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import settings
from app.files.classifier import classify_document
from app.files.models import DocumentChunk, FileVersion, ParsedDocument, ProjectFile
from app.files.ocr import (
    OCRProcessingError,
    OCRUnavailableError,
    ocr_image_bytes,
    ocr_pdf_pages,
)

MAX_CHUNK_CHARS = 4_000
CHUNKING_VERSION = 2
OCR_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png"}
MPP_SCHEDULE_TERMS = (
    "afati",
    "punim",
    "proces",
    "verbal",
    "akt",
    "njoftim",
    "fillim",
    "perfundim",
    "përfundim",
    "theme",
    "karabina",
    "fasad",
    "sistem",
    "kolaudim",
    "task",
    "duration",
    "start",
    "finish",
    "predecessor",
)
MPP_DATE_PATTERN = re.compile(
    r"\b(?:mon|tue|wed|thu|fri|sat|sun)?\s*"
    r"(?:[0-3]?\d[./-][01]?\d[./-](?:19|20)?\d{2}|"
    r"(?:jan|feb|mar|apr|may|jun|jul|aug|sep|oct|nov|dec)\s+\d{1,2},?\s+(?:19|20)?\d{2})\b",
    re.I,
)


async def parse_file_version(session: AsyncSession, *, file_version_id: uuid.UUID) -> None:
    file_version = await session.get(FileVersion, file_version_id)
    if file_version is None:
        return

    file_version.parse_status = "processing"
    await session.commit()

    suffix = Path(file_version.original_filename).suffix.lower()
    try:
        file_bytes = _load_file_version_bytes(file_version)
        if suffix == ".pdf":
            parsed = _parse_pdf(file_bytes)
        elif suffix == ".docx":
            parsed = _parse_docx(file_bytes)
        elif suffix == ".mpp":
            parsed = _parse_mpp(file_bytes)
        elif suffix in OCR_IMAGE_EXTENSIONS:
            parsed = _parse_image(file_bytes, suffix=suffix)
        else:
            file_version.parse_status = "unsupported"
            await session.commit()
            return

        completion_status = _completed_parse_status(
            suffix=suffix,
            chunks=parsed["chunks"],
            page_count=parsed["page_count"],
        )
        parsed["metadata"] = {
            **parsed["metadata"],
            "extraction_status": completion_status,
        }
        await _save_parsed_document(
            session,
            file_version=file_version,
            text_content=parsed["text_content"],
            page_count=parsed["page_count"],
            metadata=parsed["metadata"],
            chunks=parsed["chunks"],
        )
        file_version.parse_status = completion_status
        await session.commit()
    except Exception:
        await session.rollback()
        failed_version = await session.get(FileVersion, file_version_id)
        if failed_version is not None:
            failed_version.parse_status = "failed"
            await session.commit()
        raise


def _completed_parse_status(
    *,
    suffix: str,
    chunks: list[dict[str, Any]],
    page_count: int | None,
) -> str:
    readable_chunks = [chunk for chunk in chunks if str(chunk.get("text") or "").strip()]
    if any(
        dict(chunk.get("metadata") or {}).get("extraction_method") == "tesseract_tsv"
        for chunk in readable_chunks
    ):
        return "parsed_with_ocr"
    if readable_chunks:
        return "parsed"
    if suffix in {".pdf", *OCR_IMAGE_EXTENSIONS} and bool(page_count):
        return "needs_ocr"
    return "empty"


def _load_file_version_bytes(file_version: FileVersion) -> bytes:
    from app.files.storage import get_minio_client

    client = get_minio_client()
    response = None
    try:
        response = client.get_object(file_version.storage_bucket, file_version.storage_path)
        return response.read()
    finally:
        if response is not None:
            response.close()
            response.release_conn()


def _parse_pdf(pdf_bytes: bytes, *, ocr_enabled: bool | None = None) -> dict:
    reader = PdfReader(BytesIO(pdf_bytes))
    page_texts: dict[int, str] = {}
    chunks: list[dict[str, Any]] = []
    native_pages_with_text = 0
    textless_page_numbers: list[int] = []

    for index, page in enumerate(reader.pages, start=1):
        text = page.extract_text() or ""
        text = text.strip()
        if text:
            native_pages_with_text += 1
            page_texts[index] = text
            parts = _split_text(text)
            for part_index, part in enumerate(parts, start=1):
                chunks.append(
                    {
                        "text": part,
                        "page_start": index,
                        "page_end": index,
                        "metadata": {
                            "source_type": "pdf_page",
                            "extraction_method": "pypdf",
                            "page_number": index,
                            "part_index": part_index,
                            "part_count": len(parts),
                        },
                    }
                )
        else:
            textless_page_numbers.append(index)

    use_ocr = settings.ocr_enabled if ocr_enabled is None else ocr_enabled
    ocr_metadata: dict[str, Any] = {
        "enabled": use_ocr,
        "pages_requested": len(textless_page_numbers),
        "pages_with_text": 0,
        "status": "not_required" if not textless_page_numbers else "not_run",
    }
    if textless_page_numbers and use_ocr:
        try:
            ocr_result = ocr_pdf_pages(
                pdf_bytes,
                page_numbers=textless_page_numbers,
                languages=settings.ocr_languages,
                dpi=settings.ocr_dpi,
                min_confidence=settings.ocr_min_confidence,
                timeout_seconds=settings.ocr_page_timeout_seconds,
                max_chunk_chars=MAX_CHUNK_CHARS,
            )
            chunks.extend(ocr_result["chunks"])
            for page in ocr_result["pages"]:
                if str(page["text"]).strip():
                    page_texts[int(page["page_number"])] = str(page["text"])
            ocr_metadata = {
                key: value for key, value in ocr_result.items() if key not in {"pages", "chunks"}
            }
            ocr_metadata["status"] = (
                "partial" if ocr_result.get("page_errors") else "completed"
            )
            ocr_metadata["min_confidence"] = settings.ocr_min_confidence
        except (OCRUnavailableError, OCRProcessingError) as exc:
            ocr_metadata.update(
                {
                    "status": "unavailable"
                    if isinstance(exc, OCRUnavailableError)
                    else "failed",
                    "reason": str(exc)[:1_000],
                    "languages": settings.ocr_languages,
                    "dpi": settings.ocr_dpi,
                    "min_confidence": settings.ocr_min_confidence,
                }
            )
    elif textless_page_numbers:
        ocr_metadata["status"] = "disabled"

    chunks.sort(
        key=lambda chunk: (
            int(dict(chunk.get("metadata") or {}).get("page_number") or 0),
            int(dict(chunk.get("metadata") or {}).get("part_index") or 0),
        )
    )
    ocr_pages_with_text = int(ocr_metadata.get("pages_with_text") or 0)

    return {
        "text_content": "\n\n".join(
            f"--- Faqe {page_number} ---\n{page_texts[page_number]}"
            for page_number in sorted(page_texts)
        ),
        "page_count": len(reader.pages),
        "chunks": chunks,
        "metadata": {
            "parser": "pypdf+tesseract" if ocr_pages_with_text else "pypdf",
            "pages_with_text": len(page_texts),
            "native_pages_with_text": native_pages_with_text,
            "ocr_pages_with_text": ocr_pages_with_text,
            "textless_pages_after_extraction": sorted(
                set(range(1, len(reader.pages) + 1)) - set(page_texts)
            ),
            "ocr": ocr_metadata,
            "chunk_count": len(chunks),
            "chunking_version": CHUNKING_VERSION,
        },
    }


def _parse_image(
    image_bytes: bytes,
    *,
    suffix: str,
    ocr_enabled: bool | None = None,
) -> dict:
    use_ocr = settings.ocr_enabled if ocr_enabled is None else ocr_enabled
    ocr_metadata: dict[str, Any] = {
        "enabled": use_ocr,
        "pages_requested": 1,
        "pages_with_text": 0,
        "status": "not_run",
    }
    chunks: list[dict[str, Any]] = []
    text_content = ""

    if use_ocr:
        try:
            result = ocr_image_bytes(
                image_bytes,
                suffix=suffix,
                languages=settings.ocr_languages,
                dpi=settings.ocr_dpi,
                min_confidence=settings.ocr_min_confidence,
                timeout_seconds=settings.ocr_page_timeout_seconds,
                max_chunk_chars=MAX_CHUNK_CHARS,
            )
            chunks = result["chunks"]
            for chunk in chunks:
                chunk["metadata"] = {
                    **dict(chunk.get("metadata") or {}),
                    "source_type": "image_ocr",
                }
            pages = result["pages"]
            if pages and str(pages[0]["text"]).strip():
                text_content = f"--- Faqe 1 ---\n{pages[0]['text']}"
            ocr_metadata = {
                key: value for key, value in result.items() if key not in {"pages", "chunks"}
            }
            ocr_metadata["status"] = "completed"
            ocr_metadata["min_confidence"] = settings.ocr_min_confidence
        except (OCRUnavailableError, OCRProcessingError) as exc:
            ocr_metadata.update(
                {
                    "status": "unavailable"
                    if isinstance(exc, OCRUnavailableError)
                    else "failed",
                    "reason": str(exc)[:1_000],
                    "languages": settings.ocr_languages,
                    "dpi": settings.ocr_dpi,
                    "min_confidence": settings.ocr_min_confidence,
                }
            )
    else:
        ocr_metadata["status"] = "disabled"

    return {
        "text_content": text_content,
        "page_count": 1,
        "chunks": chunks,
        "metadata": {
            "parser": "tesseract",
            "pages_with_text": 1 if text_content else 0,
            "native_pages_with_text": 0,
            "ocr_pages_with_text": 1 if text_content else 0,
            "ocr": ocr_metadata,
            "chunk_count": len(chunks),
            "chunking_version": CHUNKING_VERSION,
        },
    }


def _parse_docx(docx_bytes: bytes) -> dict:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    document = Document(BytesIO(docx_bytes))
    text_blocks: list[str] = []
    chunks: list[dict[str, Any]] = []
    pending_paragraphs: list[dict[str, Any]] = []
    paragraph_index = 0
    table_index = 0
    block_index = 0

    for block in document.iter_inner_content():
        block_index += 1
        if isinstance(block, Paragraph):
            paragraph_index += 1
            text = block.text.strip()
            if not text:
                continue
            text_blocks.append(text)
            pending_paragraphs.append(
                {
                    "text": text,
                    "paragraph_index": paragraph_index,
                    "block_index": block_index,
                }
            )
            continue

        if not isinstance(block, Table):
            continue

        chunks.extend(_docx_paragraph_chunks(pending_paragraphs))
        pending_paragraphs = []
        table_index += 1
        table_rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(block.rows, start=1):
            cells = [_normalize_cell_text(cell.text) for cell in row.cells]
            cells = [cell for cell in cells if cell]
            if cells:
                table_rows.append(
                    {
                        "text": " | ".join(cells),
                        "row_index": row_index,
                    }
                )
        if table_rows:
            text_blocks.append(
                f"--- Tabela {table_index} ---\n"
                + "\n".join(str(row["text"]) for row in table_rows)
            )
            chunks.extend(
                _docx_table_chunks(
                    table_rows,
                    table_index=table_index,
                    block_index=block_index,
                )
            )

    chunks.extend(_docx_paragraph_chunks(pending_paragraphs))

    return {
        "text_content": "\n\n".join(text_blocks),
        "page_count": None,
        "chunks": chunks,
        "metadata": {
            "parser": "python-docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
            "text_blocks": len(text_blocks),
            "chunk_count": len(chunks),
            "chunking_version": CHUNKING_VERSION,
        },
    }


def _parse_mpp(
    mpp_bytes: bytes,
    *,
    extractor_command: str | None = None,
) -> dict:
    command = settings.mpp_extractor_command if extractor_command is None else extractor_command
    external_error: str | None = None
    if command:
        try:
            parsed = _parse_mpp_with_command(mpp_bytes, command=command)
            if parsed["chunks"]:
                return parsed
        except Exception as exc:
            external_error = str(exc)[:1_000]

    parsed = _parse_mpp_best_effort_strings(mpp_bytes)
    parsed["metadata"]["external_extractor"] = {
        "configured": bool(command),
        "status": "failed" if external_error else "not_configured",
        "error": external_error,
    }
    return parsed


def _parse_mpp_with_command(mpp_bytes: bytes, *, command: str) -> dict:
    with TemporaryDirectory(prefix="atf-mpp-") as temp_dir:
        input_path = Path(temp_dir) / "input.mpp"
        output_path = Path(temp_dir) / "output.json"
        input_path.write_bytes(mpp_bytes)

        formatted_command = command.format(
            input=str(input_path),
            output=str(output_path),
        )
        result = subprocess.run(  # noqa: S603 - operator-configured local extractor.
            shlex.split(formatted_command),
            check=False,
            capture_output=True,
            timeout=settings.mpp_extractor_timeout_seconds,
        )
        if result.returncode != 0:
            stderr = result.stderr.decode("utf-8", errors="replace")
            stdout = result.stdout.decode("utf-8", errors="replace")
            raise RuntimeError(
                f"MPP extractor failed with code {result.returncode}: "
                f"{(stderr or stdout)[:500]}"
            )

        output_text = ""
        if output_path.exists() and output_path.stat().st_size:
            output_text = output_path.read_text(encoding="utf-8", errors="replace")
        if not output_text:
            output_text = result.stdout.decode("utf-8", errors="replace")
        if not output_text.strip():
            raise RuntimeError("MPP extractor returned no output.")

    return _parse_mpp_external_output(output_text)


def _parse_mpp_external_output(output_text: str) -> dict:
    try:
        payload = json.loads(output_text)
    except json.JSONDecodeError:
        return _mpp_text_lines_to_parsed(
            _clean_extracted_lines(output_text.splitlines()),
            parser="mpp-external-text",
            extraction_method="external_command_text",
            extraction_status="parsed",
            extra_metadata={"structured": False},
        )

    tasks = _mpp_tasks_from_payload(payload)
    if tasks:
        lines = [_mpp_task_line(index, task) for index, task in enumerate(tasks, start=1)]
        return _mpp_text_lines_to_parsed(
            lines,
            parser="mpp-external-json",
            extraction_method="external_command_json",
            extraction_status="parsed",
            extra_metadata={
                "structured": True,
                "task_count": len(tasks),
            },
        )

    if isinstance(payload, dict):
        text = json.dumps(payload, ensure_ascii=False, indent=2)
    else:
        text = json.dumps({"data": payload}, ensure_ascii=False, indent=2)
    return _mpp_text_lines_to_parsed(
        _clean_extracted_lines(text.splitlines()),
        parser="mpp-external-json",
        extraction_method="external_command_json",
        extraction_status="partial",
        extra_metadata={"structured": False},
    )


def _mpp_tasks_from_payload(payload: Any) -> list[dict[str, Any]]:
    if isinstance(payload, list):
        return [item for item in payload if isinstance(item, dict)]
    if not isinstance(payload, dict):
        return []
    for key in ("tasks", "activities", "rows", "data"):
        value = payload.get(key)
        if isinstance(value, list):
            return [item for item in value if isinstance(item, dict)]
    return []


def _mpp_task_line(index: int, task: dict[str, Any]) -> str:
    field_aliases = {
        "id": ("id", "uid", "unique_id", "uniqueID"),
        "wbs": ("wbs", "outline_number", "outlineNumber"),
        "name": ("name", "task_name", "taskName", "text", "title"),
        "duration": ("duration", "duration_text", "durationText"),
        "start": ("start", "start_date", "startDate"),
        "finish": ("finish", "finish_date", "finishDate", "end"),
        "predecessors": ("predecessors", "predecessor", "predecessor_ids"),
        "percent_complete": ("percent_complete", "percentComplete", "complete"),
        "resources": ("resources", "resource_names", "resourceNames"),
    }
    parts = [f"Rreshti {index}"]
    for label, aliases in field_aliases.items():
        value = next(
            (
                task.get(alias)
                for alias in aliases
                if task.get(alias) not in {None, "", []}
            ),
            None,
        )
        if isinstance(value, list):
            value = ", ".join(str(item) for item in value if str(item).strip())
        if value is not None and str(value).strip():
            parts.append(f"{label}: {value}")
    return " | ".join(parts)


def _parse_mpp_best_effort_strings(mpp_bytes: bytes) -> dict:
    ascii_strings = _extract_ascii_strings(mpp_bytes)
    utf16_strings = _extract_utf16le_strings(mpp_bytes)
    raw_lines = _deduplicate_lines([*utf16_strings, *ascii_strings])
    candidate_lines = _select_mpp_schedule_lines(raw_lines)
    if not candidate_lines:
        candidate_lines = raw_lines[:300]

    heading = (
        "Ekstrakt tekstual best-effort nga skedari Microsoft Project (MPP). "
        "Përdoret për kronologji orientuese; varësitë, datat dhe përqindjet duhen "
        "verifikuar me skedarin origjinal kur mungojnë në tekst."
    )
    lines = [heading, *candidate_lines]
    return _mpp_text_lines_to_parsed(
        lines,
        parser="mpp-best-effort-strings",
        extraction_method="binary_string_scan",
        extraction_status="partial",
        extra_metadata={
            "ascii_string_count": len(ascii_strings),
            "utf16_string_count": len(utf16_strings),
            "selected_line_count": len(candidate_lines),
        },
    )


def _mpp_text_lines_to_parsed(
    lines: list[str],
    *,
    parser: str,
    extraction_method: str,
    extraction_status: str,
    extra_metadata: dict[str, Any] | None = None,
) -> dict:
    clean_lines = _clean_extracted_lines(lines)
    text_content = "\n".join(clean_lines)
    chunks: list[dict[str, Any]] = []
    current: list[tuple[int, str]] = []
    current_length = 0

    def flush() -> None:
        nonlocal current, current_length
        if not current:
            return
        chunks.append(
            {
                "text": "\n".join(line for _, line in current),
                "page_start": None,
                "page_end": None,
                "metadata": {
                    "source_type": "mpp_schedule",
                    "extraction_method": extraction_method,
                    "row_start": current[0][0],
                    "row_end": current[-1][0],
                },
            }
        )
        current = []
        current_length = 0

    for row_index, line in enumerate(clean_lines, start=1):
        line_parts = _split_text(line)
        for part in line_parts:
            separator_length = 1 if current else 0
            if current and current_length + separator_length + len(part) > MAX_CHUNK_CHARS:
                flush()
            current.append((row_index, part))
            current_length += (1 if current_length else 0) + len(part)
    flush()

    return {
        "text_content": text_content,
        "page_count": None,
        "chunks": chunks,
        "metadata": {
            "parser": parser,
            "extraction_status": extraction_status,
            "line_count": len(clean_lines),
            "chunk_count": len(chunks),
            "chunking_version": CHUNKING_VERSION,
            **(extra_metadata or {}),
        },
    }


def _extract_ascii_strings(data: bytes, *, min_length: int = 4) -> list[str]:
    strings: list[str] = []
    current = bytearray()
    for byte in data:
        if byte in {9, 10, 13} or 32 <= byte <= 126:
            current.append(byte)
            continue
        if len(current) >= min_length:
            strings.append(current.decode("latin-1", errors="ignore"))
        current = bytearray()
    if len(current) >= min_length:
        strings.append(current.decode("latin-1", errors="ignore"))
    return strings


def _extract_utf16le_strings(data: bytes, *, min_length: int = 4) -> list[str]:
    strings: list[str] = []
    for offset in (0, 1):
        current: list[str] = []
        for index in range(offset, len(data) - 1, 2):
            codepoint = data[index] | (data[index + 1] << 8)
            character = chr(codepoint)
            if character in {"\t", "\n", "\r"} or (
                character.isprintable() and not character.isspace()
            ) or character == " ":
                current.append(character)
                continue
            if len(current) >= min_length:
                strings.append("".join(current))
            current = []
        if len(current) >= min_length:
            strings.append("".join(current))
    return strings


def _select_mpp_schedule_lines(lines: list[str]) -> list[str]:
    selected: list[str] = []
    for line in lines:
        normalized = _normalize_mpp_line(line).lower()
        has_term = any(term in normalized for term in MPP_SCHEDULE_TERMS)
        has_date = bool(MPP_DATE_PATTERN.search(normalized))
        if has_term or has_date:
            selected.append(line)
    return selected[:500]


def _clean_extracted_lines(lines: list[str]) -> list[str]:
    return _deduplicate_lines(
        [
            cleaned
            for line in lines
            if (cleaned := _normalize_mpp_line(str(line)))
            and _is_useful_mpp_line(cleaned)
        ]
    )


def _normalize_mpp_line(line: str) -> str:
    line = line.replace("\x00", " ")
    line = re.sub(r"[\x01-\x08\x0b-\x1f\x7f]+", " ", line)
    line = " ".join(line.split())
    return line.strip()


def _is_useful_mpp_line(line: str) -> bool:
    if len(line) < 3 or len(line) > 500:
        return False
    lowered = line.lower()
    low_value_markers = {
        "root entry",
        "compobj",
        "summaryinformation",
        "document summaryinformation",
        "microsoft project",
        "projectdata",
    }
    if lowered in low_value_markers:
        return False
    alpha_count = sum(1 for character in line if character.isalpha())
    digit_count = sum(1 for character in line if character.isdigit())
    return alpha_count >= 2 or digit_count >= 4


def _deduplicate_lines(lines: list[str]) -> list[str]:
    seen: set[str] = set()
    deduplicated: list[str] = []
    for line in lines:
        key = line.casefold()
        if key in seen:
            continue
        seen.add(key)
        deduplicated.append(line)
    return deduplicated


def _normalize_cell_text(value: str) -> str:
    return " ".join(value.split())


def _split_text(text: str, *, max_chars: int = MAX_CHUNK_CHARS) -> list[str]:
    text = text.strip()
    if not text:
        return []

    blocks = [block.strip() for block in text.splitlines() if block.strip()]
    if not blocks:
        return []

    parts: list[str] = []
    current: list[str] = []
    current_length = 0
    for block in blocks:
        block_parts = _split_oversized_block(block, max_chars=max_chars)
        for block_part in block_parts:
            separator_length = 1 if current else 0
            if current and current_length + separator_length + len(block_part) > max_chars:
                parts.append("\n".join(current))
                current = []
                current_length = 0
            current.append(block_part)
            current_length += (1 if current_length else 0) + len(block_part)

    if current:
        parts.append("\n".join(current))
    return parts


def _split_oversized_block(block: str, *, max_chars: int) -> list[str]:
    if len(block) <= max_chars:
        return [block]

    parts: list[str] = []
    remaining = block
    while len(remaining) > max_chars:
        split_at = remaining.rfind(" ", 0, max_chars + 1)
        if split_at <= 0:
            split_at = max_chars
        parts.append(remaining[:split_at].strip())
        remaining = remaining[split_at:].strip()
    if remaining:
        parts.append(remaining)
    return parts


def _docx_paragraph_chunks(paragraphs: list[dict[str, Any]]) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []
    current_length = 0

    def flush() -> None:
        nonlocal current, current_length
        if not current:
            return
        chunks.append(
            {
                "text": "\n\n".join(str(item["text"]) for item in current),
                "page_start": None,
                "page_end": None,
                "metadata": {
                    "source_type": "docx_paragraphs",
                    "paragraph_start": current[0]["paragraph_index"],
                    "paragraph_end": current[-1]["paragraph_index"],
                    "block_start": current[0]["block_index"],
                    "block_end": current[-1]["block_index"],
                },
            }
        )
        current = []
        current_length = 0

    for paragraph in paragraphs:
        paragraph_parts = _split_text(str(paragraph["text"]))
        for part_index, part in enumerate(paragraph_parts, start=1):
            item = {
                **paragraph,
                "text": part,
                "part_index": part_index,
                "part_count": len(paragraph_parts),
            }
            separator_length = 2 if current else 0
            if current and current_length + separator_length + len(part) > MAX_CHUNK_CHARS:
                flush()
            current.append(item)
            current_length += (2 if current_length else 0) + len(part)
    flush()
    return chunks


def _docx_table_chunks(
    rows: list[dict[str, Any]],
    *,
    table_index: int,
    block_index: int,
) -> list[dict[str, Any]]:
    chunks: list[dict[str, Any]] = []
    current: list[dict[str, Any]] = []
    current_length = 0

    def flush() -> None:
        nonlocal current, current_length
        if not current:
            return
        chunks.append(
            {
                "text": "\n".join(str(item["text"]) for item in current),
                "page_start": None,
                "page_end": None,
                "metadata": {
                    "source_type": "docx_table",
                    "table_index": table_index,
                    "row_start": current[0]["row_index"],
                    "row_end": current[-1]["row_index"],
                    "block_index": block_index,
                },
            }
        )
        current = []
        current_length = 0

    for row in rows:
        row_parts = _split_text(str(row["text"]))
        for part in row_parts:
            item = {**row, "text": part}
            separator_length = 1 if current else 0
            if current and current_length + separator_length + len(part) > MAX_CHUNK_CHARS:
                flush()
            current.append(item)
            current_length += (1 if current_length else 0) + len(part)
    flush()
    return chunks


async def _save_parsed_document(
    session: AsyncSession,
    *,
    file_version: FileVersion,
    text_content: str,
    page_count: int | None,
    metadata: dict,
    chunks: list[dict[str, Any]],
) -> None:
    chunks = [chunk for chunk in chunks if str(chunk.get("text") or "").strip()]
    classification = classify_document(file_version.original_filename, text_content)
    parsed_metadata = {
        **metadata,
        "classification": classification.as_metadata(),
        "chunk_count": len(chunks),
        "chunking_version": CHUNKING_VERSION,
    }

    result = await session.execute(
        select(ParsedDocument).where(ParsedDocument.file_version_id == file_version.id)
    )
    parsed_document = result.scalar_one_or_none()

    if parsed_document is None:
        parsed_document = ParsedDocument(
            file_version_id=file_version.id,
            document_type=classification.document_type,
            language="sq-AL",
            page_count=page_count,
            text_content=text_content,
            document_metadata=parsed_metadata,
        )
        session.add(parsed_document)
        await session.flush()
    else:
        parsed_document.document_type = classification.document_type
        parsed_document.page_count = page_count
        parsed_document.text_content = text_content
        parsed_document.document_metadata = parsed_metadata

    project_id = await get_project_id_for_file_version(
        session,
        file_version_id=file_version.id,
    )
    if project_id is None:
        raise ValueError("File version is not linked to a project.")

    await _replace_document_chunks(
        session,
        parsed_document=parsed_document,
        file_version=file_version,
        project_id=project_id,
        chunks=chunks,
    )


async def _replace_document_chunks(
    session: AsyncSession,
    *,
    parsed_document: ParsedDocument,
    file_version: FileVersion,
    project_id: uuid.UUID,
    chunks: list[dict[str, Any]],
) -> None:
    stored_chunks = [chunk for chunk in chunks if str(chunk.get("text") or "").strip()]
    await session.execute(
        delete(DocumentChunk).where(DocumentChunk.file_version_id == file_version.id)
    )
    session.add_all(
        [
            DocumentChunk(
                parsed_document_id=parsed_document.id,
                file_version_id=file_version.id,
                project_id=project_id,
                chunk_index=index,
                page_start=chunk.get("page_start"),
                page_end=chunk.get("page_end"),
                text=str(chunk["text"]),
                chunk_metadata={
                    **dict(chunk.get("metadata") or {}),
                    "chunking_version": CHUNKING_VERSION,
                },
            )
            for index, chunk in enumerate(stored_chunks)
        ]
    )


async def get_project_id_for_file_version(
    session: AsyncSession,
    *,
    file_version_id: uuid.UUID,
) -> uuid.UUID | None:
    result = await session.execute(
        select(ProjectFile.project_id)
        .join(FileVersion, FileVersion.file_id == ProjectFile.id)
        .where(FileVersion.id == file_version_id)
    )
    return result.scalar_one_or_none()
